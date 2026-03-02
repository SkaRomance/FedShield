import argparse
import math
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from PIL import Image, ImageDraw


def iter_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph


def replace_placeholder_text(paragraph, placeholder: str, value: str) -> bool:
    replaced = False
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            replaced = True
    if replaced:
        return True
    if placeholder not in paragraph.text:
        return False
    merged = paragraph.text.replace(placeholder, value)
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = merged
    else:
        paragraph.add_run(merged)
    return True


def _star_points(cx: float, cy: float, outer: float, inner: float):
    points = []
    for index in range(10):
        angle_deg = -90 + index * 36
        angle_rad = angle_deg * 3.141592653589793 / 180
        radius = outer if index % 2 == 0 else inner
        x = cx + radius * math.cos(angle_rad)
        y = cy + radius * math.sin(angle_rad)
        points.append((x, y))
    return points


def generate_stars_png(target_path: Path, stars: int):
    clamped = max(0, min(5, int(round(stars))))
    if clamped == 0:
        img = Image.new("RGBA", (32, 32), (255, 255, 255, 0))
        img.save(target_path)
        return

    star_size = 44
    gap = 8
    padding = 6
    width = clamped * star_size + (clamped - 1) * gap + padding * 2
    height = star_size + padding * 2
    img = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    orange = (239, 89, 0, 255)

    for idx in range(clamped):
        left = padding + idx * (star_size + gap)
        cx = left + star_size / 2
        cy = padding + star_size / 2
        points = _star_points(cx, cy, outer=star_size * 0.48, inner=star_size * 0.2)
        draw.polygon(points, fill=orange, outline=orange)

    img.save(target_path)


def replace_placeholder_image(paragraph, placeholder: str, image_path: Path, width_pt: float, force_center: bool = False) -> bool:
    if placeholder not in paragraph.text:
        return False
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Pt(width_pt))
    if force_center:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return True


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--output-docx", required=True)
    parser.add_argument("--company-name", required=True)
    parser.add_argument("--partita-iva", required=True)
    parser.add_argument("--data-sopralluogo", required=True)
    parser.add_argument("--fed-score", required=True)
    parser.add_argument("--fed-stars", required=True, type=int)
    parser.add_argument("--work-dir", required=True)
    parser.add_argument("--qr-image", required=False)
    args = parser.parse_args()

    template_path = Path(args.template)
    output_path = Path(args.output_docx)
    work_dir = Path(args.work_dir)
    work_dir.mkdir(parents=True, exist_ok=True)

    document = Document(str(template_path))

    placeholders_found = {
        "{{RAGIONE_SOCIALE}}": False,
        "{{PARTITA_IVA}}": False,
        "{{DATA_SOPRALLUOGO}}": False,
        "{{FED_SCORE}}": False,
        "{{FED_STARS}}": False,
        "{{QR_CODE}}": False,
    }

    stars_image_path = work_dir / "fed_stars.png"
    generate_stars_png(stars_image_path, args.fed_stars)

    for paragraph in iter_paragraphs(document):
        if replace_placeholder_text(paragraph, "{{RAGIONE_SOCIALE}}", args.company_name):
            placeholders_found["{{RAGIONE_SOCIALE}}"] = True
        if replace_placeholder_text(paragraph, "{{PARTITA_IVA}}", args.partita_iva):
            placeholders_found["{{PARTITA_IVA}}"] = True
        if replace_placeholder_text(paragraph, "{{DATA_SOPRALLUOGO}}", args.data_sopralluogo):
            placeholders_found["{{DATA_SOPRALLUOGO}}"] = True
        if replace_placeholder_text(paragraph, "{{FED_SCORE}}", str(args.fed_score)):
            placeholders_found["{{FED_SCORE}}"] = True
        if replace_placeholder_image(paragraph, "{{FED_STARS}}", stars_image_path, width_pt=180, force_center=True):
            placeholders_found["{{FED_STARS}}"] = True
        if args.qr_image:
            qr_path = Path(args.qr_image)
            if replace_placeholder_image(paragraph, "{{QR_CODE}}", qr_path, width_pt=96, force_center=True):
                placeholders_found["{{QR_CODE}}"] = True
        else:
            if replace_placeholder_text(paragraph, "{{QR_CODE}}", "QR non disponibile"):
                placeholders_found["{{QR_CODE}}"] = True

    missing = [key for key, found in placeholders_found.items() if not found]
    if missing:
        raise RuntimeError(f"Placeholder non trovati nel template: {', '.join(missing)}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


if __name__ == "__main__":
    main()
